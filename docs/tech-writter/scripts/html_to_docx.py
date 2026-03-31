#!/usr/bin/env python3
"""HTML to DOCX converter with CJK font support."""
import sys
import os
import re
from html.parser import HTMLParser

def convert(html_path: str, output_path: str = None):
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx")
        sys.exit(1)

    if not os.path.exists(html_path):
        print(f"ERROR: File not found: {html_path}")
        sys.exit(1)

    if output_path is None:
        output_path = os.path.splitext(html_path)[0] + ".docx"

    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    doc = Document()

    # Set default CJK font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(11)

    # Set heading fonts
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Microsoft JhengHei'
        heading_style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    # Simple HTML text extraction for content
    class ContentExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self.sections = []
            self.current_text = []
            self.current_tag = None
            self.in_pre = False
            self.in_table = False
            self.table_rows = []
            self.current_row = []
            self.skip_tags = {'style', 'script', 'nav'}
            self.skip_depth = 0

        def handle_starttag(self, tag, attrs):
            if tag in self.skip_tags:
                self.skip_depth += 1
                return
            if self.skip_depth > 0:
                return
            if tag in ('h1', 'h2', 'h3'):
                self.flush_text()
                self.current_tag = tag
            elif tag == 'pre':
                self.flush_text()
                self.in_pre = True
            elif tag == 'table':
                self.flush_text()
                self.in_table = True
                self.table_rows = []
            elif tag == 'tr':
                self.current_row = []
            elif tag in ('p', 'div', 'li'):
                if not self.in_table:
                    self.flush_text()
                    self.current_tag = tag

        def handle_endtag(self, tag):
            if tag in self.skip_tags:
                self.skip_depth -= 1
                return
            if self.skip_depth > 0:
                return
            if tag in ('h1', 'h2', 'h3'):
                text = ''.join(self.current_text).strip()
                if text:
                    level = int(tag[1])
                    self.sections.append(('heading', level, text))
                self.current_text = []
                self.current_tag = None
            elif tag == 'pre':
                text = ''.join(self.current_text).strip()
                if text:
                    self.sections.append(('code', 0, text))
                self.current_text = []
                self.in_pre = False
            elif tag == 'table':
                if self.table_rows:
                    self.sections.append(('table', 0, self.table_rows))
                self.in_table = False
                self.table_rows = []
            elif tag == 'tr':
                if self.current_row:
                    self.table_rows.append(self.current_row)
            elif tag in ('td', 'th'):
                self.current_row.append(''.join(self.current_text).strip())
                self.current_text = []
            elif tag in ('p', 'div', 'li'):
                self.flush_text()

        def handle_data(self, data):
            if self.skip_depth > 0:
                return
            self.current_text.append(data)

        def flush_text(self):
            text = ''.join(self.current_text).strip()
            if text and not self.in_table:
                self.sections.append(('paragraph', 0, text))
            self.current_text = []

    extractor = ContentExtractor()
    extractor.feed(html_content)
    extractor.flush_text()

    for section_type, level, content in extractor.sections:
        if section_type == 'heading':
            doc.add_heading(content, level=level)
        elif section_type == 'paragraph':
            doc.add_paragraph(content)
        elif section_type == 'code':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)
        elif section_type == 'table' and content:
            cols = max(len(row) for row in content)
            table = doc.add_table(rows=len(content), cols=cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, row in enumerate(content):
                for j, cell in enumerate(row):
                    if j < cols:
                        table.rows[i].cells[j].text = cell
                        for p in table.rows[i].cells[j].paragraphs:
                            for run in p.runs:
                                run.font.name = 'Microsoft JhengHei'
                                run.font.size = Pt(10)

    doc.save(output_path)
    print(f"DOCX generated: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python html_to_docx.py <input.html> [output.docx]")
        sys.exit(1)
    out = sys.argv[2] if len(sys.argv) > 2 else None
    convert(sys.argv[1], out)
